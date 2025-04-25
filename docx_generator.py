import os
import io
import logging
import datetime
import json
import tempfile
from flask import send_file, url_for
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# Configure logging
logging.basicConfig(level=logging.DEBUG)

# Directory to store generated files temporarily
TEMP_DIR = os.path.join('static', 'downloads')
os.makedirs(TEMP_DIR, exist_ok=True)

def generate_docx(title, content):
    """
    Generate a .docx file from the given title and content
    
    Args:
        title (str): Document title
        content (str): Document content or JSON string of content
        
    Returns:
        str: URL path to the generated .docx file
    """
    try:
        # Create a new Document
        doc = Document()
        
        # Set default font and margins
        sections = doc.sections
        for section in sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
        
        # Add title with formatting
        title_heading = doc.add_heading(title, level=1)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title_heading.runs:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Add timestamp with styling
        now = datetime.datetime.now()
        timestamp_para = doc.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_run = timestamp_para.add_run(f"Generated on: {now.strftime('%Y-%m-%d %H:%M:%S')}")
        timestamp_run.font.size = Pt(10)
        timestamp_run.font.italic = True
        timestamp_run.font.color.rgb = RGBColor(100, 100, 100)  # Dark gray
        
        # Add source with styling
        source_para = doc.add_paragraph()
        source_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_run = source_para.add_run("Source: WikiTruth")
        source_run.font.size = Pt(10)
        source_run.font.italic = True
        source_run.font.color.rgb = RGBColor(100, 100, 100)  # Dark gray
        
        # Add a horizontal line/separator
        separator = doc.add_paragraph()
        separator.paragraph_format.space_before = Pt(10)
        separator.paragraph_format.space_after = Pt(10)
        separator_run = separator.add_run("_" * 50)
        separator_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray
        
        # Parse and add content
        if isinstance(content, str) and (content.startswith('{') and content.endswith('}')):
            try:
                content_dict = json.loads(content)
                if isinstance(content_dict, dict):
                    # This is a dictionary of sections
                    for section, text in content_dict.items():
                        if section == 'Introduction':
                            # Format the introduction paragraph
                            intro_para = doc.add_paragraph()
                            intro_para.paragraph_format.space_after = Pt(12)
                            intro_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                            intro_run = intro_para.add_run(text)
                            intro_run.font.size = Pt(12)
                        else:
                            # Format section headings
                            section_heading = doc.add_heading(section, level=2)
                            for run in section_heading.runs:
                                run.font.size = Pt(14)
                                run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                            
                            # Format section content with paragraph breaks
                            paragraphs = text.split('\n\n')
                            for p_text in paragraphs:
                                if p_text.strip():
                                    # Check if this is a list item
                                    if p_text.strip().startswith('•'):
                                        list_items = p_text.strip().split('\n')
                                        for item in list_items:
                                            if item.strip():
                                                list_para = doc.add_paragraph(
                                                    item.strip().lstrip('• '), 
                                                    style='List Bullet'
                                                )
                                                list_para.paragraph_format.space_after = Pt(6)
                                                for run in list_para.runs:
                                                    run.font.size = Pt(12)
                                    else:
                                        section_para = doc.add_paragraph()
                                        section_para.paragraph_format.space_after = Pt(12)
                                        section_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                                        section_run = section_para.add_run(p_text.strip())
                                        section_run.font.size = Pt(12)
                else:
                    # It's JSON but not a dictionary - use as plain text
                    para = doc.add_paragraph()
                    para.paragraph_format.space_after = Pt(12)
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    para_run = para.add_run(content)
                    para_run.font.size = Pt(12)
            except json.JSONDecodeError:
                # Not valid JSON, use as plain text
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(12)
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                para_run = para.add_run(content)
                para_run.font.size = Pt(12)
        elif isinstance(content, dict):
            # Direct dictionary input (not JSON string)
            for section, text in content.items():
                if section == 'Introduction':
                    # Format the introduction paragraph
                    intro_para = doc.add_paragraph()
                    intro_para.paragraph_format.space_after = Pt(12)
                    intro_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    intro_run = intro_para.add_run(text)
                    intro_run.font.size = Pt(12)
                else:
                    # Format section headings
                    section_heading = doc.add_heading(section, level=2)
                    for run in section_heading.runs:
                        run.font.size = Pt(14)
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                    
                    # Format section content
                    section_para = doc.add_paragraph()
                    section_para.paragraph_format.space_after = Pt(12)
                    section_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    section_run = section_para.add_run(text)
                    section_run.font.size = Pt(12)
        else:
            # Simple string content (e.g., for summary)
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para_run = para.add_run(content)
            para_run.font.size = Pt(12)
        
        # Add footer with page numbers
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Page ")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        
        # Create a safe filename
        safe_title = "".join([c for c in title if c.isalnum() or c in ' .-_']).rstrip()
        if not safe_title:
            safe_title = "document"
        timestamp = now.strftime('%Y%m%d%H%M%S')
        filename = f"{safe_title}_{timestamp}.docx"
        
        # Ensure downloads directory exists
        os.makedirs(TEMP_DIR, exist_ok=True)
        
        # Save to a physical file in the static/downloads directory
        file_path = os.path.join(TEMP_DIR, filename)
        doc.save(file_path)
        
        # Return the URL to the file
        return url_for('static', filename=f'downloads/{filename}')
    
    except Exception as e:
        logging.error(f"DOCX generation error: {str(e)}")
        raise Exception(f"Failed to generate DOCX file: {str(e)}")
