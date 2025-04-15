import os
import json
import logging
import time
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file
from docx import Document
from utils import get_language_name, timestamp_to_date
import wiki_utils
import threading
import requests


# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

# Create the Flask application
app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET", "wikitruth_secret_key")

# Data storage (file-based)
HIGHLIGHTS_FILE = "highlights.json"

# Common language list
LANGUAGES = {
    'en': 'English',
    'es': 'Spanish',
    'fr': 'French',
    'de': 'German',
    'it': 'Italian',
    'pt': 'Portuguese',
    'zh': 'Chinese',
    'ja': 'Japanese',
    'ru': 'Russian',
    'ar': 'Arabic',
    'hi': 'Hindi',
    'ko': 'Korean',
    'tr': 'Turkish',
}

# Template filters
@app.template_filter('timestamp_to_date')
def _timestamp_to_date(timestamp):
    """Convert timestamp to formatted date"""
    return timestamp_to_date(timestamp)

# Ensure highlights file exists
def ensure_highlights_file():
    if not os.path.exists(HIGHLIGHTS_FILE):
        with open(HIGHLIGHTS_FILE, 'w') as f:
            json.dump({}, f)

# Load highlights from file
def load_highlights():
    ensure_highlights_file()
    try:
        with open(HIGHLIGHTS_FILE, 'r') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Error loading highlights: {str(e)}")
        return {}

# Save highlights to file
def save_highlights(highlights_data):
    ensure_highlights_file()
    try:
        with open(HIGHLIGHTS_FILE, 'w') as f:
            json.dump(highlights_data, f)
        return True
    except Exception as e:
        logger.error(f"Error saving highlights: {str(e)}")
        return False

# Template context processors
@app.context_processor
def inject_globals():
    """Inject global variables into templates"""
    return {
        'language_dict': LANGUAGES,
        'get_language_name': get_language_name
    }

# Routes
@app.route('/')
def home():
    """Home page with search form"""
    selected_lang = session.get('search_lang', 'en')
    return render_template('home.html', selected_lang=selected_lang)

@app.route('/search', methods=['GET', 'POST'])
def search():
    """Search Wikipedia for articles"""
    if request.method == 'POST':
        search_query = request.form.get('search_query', '')
        search_lang = request.form.get('search_lang', 'en')
        
        # Store in session for easy access
        session['search_query'] = search_query
        session['search_lang'] = search_lang
        
        if not search_query:
            flash('Please enter a search term', 'danger')
            return redirect(url_for('home'))
        
        # Search Wikipedia
        try:
            search_results = wiki_utils.get_wikipedia_search_results(search_query, search_lang)
            return render_template('search_results.html', 
                                  search_query=search_query,
                                  search_lang=search_lang, 
                                  search_results=search_results)
        except Exception as e:
            logger.error(f"Search error: {str(e)}")
            flash(f"Error searching Wikipedia: {str(e)}", 'danger')
            return redirect(url_for('home'))
    else:
        # GET request - redirect to home
        return redirect(url_for('home'))

@app.route('/article/<title>')
def view_article(title):
    """View a Wikipedia article"""
    lang = request.args.get('lang', 'en')
    show_translation = request.args.get('translate', 'false') == 'true'
    translate_to = request.args.get('to_lang', 'en')
    
    try:
        # Get article content from Wikipedia
        article = wiki_utils.get_article_content(title, lang)
        
        if not article:
            flash('Article not found', 'danger')
            return redirect(url_for('home'))
        
        # Process article content into sections
        sections = wiki_utils.split_content_into_sections(article['content'])
        
        # Get available languages for this article
        available_languages = wiki_utils.get_available_languages(title, lang)
        
        return render_template('article.html', 
                              article=article,
                              sections=sections,
                              lang=lang,
                              show_translation=show_translation,
                              translate_to=translate_to,
                              available_languages=available_languages)
    except Exception as e:
        logger.error(f"Article view error: {str(e)}")
        flash(f"Error retrieving article: {str(e)}", 'danger')
        return redirect(url_for('home'))

@app.route('/translate-section', methods=['POST'])
def translate_section():
    """Translate a section of text"""
    text = request.form.get('text', '')
    from_lang = request.form.get('from_lang', 'auto')
    to_lang = request.form.get('to_lang', 'en')
    
    if not text or not to_lang:
        return jsonify({'error': 'Missing required parameters'})
    
    try:
        translated_text = wiki_utils.translate_text(text, to_lang, from_lang)
        return jsonify({'translated_text': translated_text})
    except Exception as e:
        logger.error(f"Translation error: {str(e)}")
        return jsonify({'error': f'Translation error: {str(e)}'})

@app.route('/save-highlight', methods=['POST'])
def save_highlight():
    """Save a highlighted text for review"""
    article_id = request.form.get('article_id', '')
    text_to_highlight = request.form.get('text_to_highlight', '')
    context = request.form.get('context', '')
    
    # Check if this is just a retrieval request
    if article_id and context == 'retrieve_only':
        try:
            # Load existing highlights
            highlights_data = load_highlights()
            
            # Parse article_id (format: title_lang)
            try:
                title, lang = article_id.rsplit('_', 1)
            except ValueError:
                title = article_id
                lang = 'en'
            
            # Get highlights for this article
            if title in highlights_data and lang in highlights_data[title]['languages']:
                article_highlights = highlights_data[title]['languages'][lang]['highlights']
                return jsonify({'success': True, 'highlights': article_highlights})
            else:
                return jsonify({'success': True, 'highlights': []})
        
        except Exception as e:
            logger.error(f"Error retrieving highlights: {str(e)}")
            return jsonify({'success': False, 'error': str(e)})
    
    # If not a retrieval request, require text_to_highlight
    if not article_id or not text_to_highlight:
        return jsonify({'success': False, 'error': 'Missing required parameters'})
    
    try:
        # Load existing highlights
        highlights_data = load_highlights()
        
        # Parse article_id (format: title_lang)
        try:
            title, lang = article_id.rsplit('_', 1)
        except ValueError:
            title = article_id
            lang = 'en'
        
        # Create structure if it doesn't exist
        if title not in highlights_data:
            highlights_data[title] = {'languages': {}}
        
        if lang not in highlights_data[title]['languages']:
            highlights_data[title]['languages'][lang] = {
                'language_name': get_language_name(lang),
                'highlights': []
            }
        
        # Add the highlight
        highlights_data[title]['languages'][lang]['highlights'].append({
            'text': text_to_highlight,
            'context': context,
            'timestamp': int(time.time())
        })
        
        # Save highlights
        if save_highlights(highlights_data):
            # Return the updated highlights
            article_highlights = highlights_data[title]['languages'][lang]['highlights']
            return jsonify({'success': True, 'highlights': article_highlights})
        else:
            return jsonify({'success': False, 'error': 'Error saving highlight'})
    except Exception as e:
        logger.error(f"Error saving highlight: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/highlights')
def view_all_highlights():
    """View all highlights/reviews"""
    highlights_data = load_highlights()
    return render_template('highlights.html', articles=highlights_data)

@app.route('/export/<title>')
def export_article(title):
    """Export article as Word document"""
    lang = request.args.get('lang', 'en')
    include_translations = request.args.get('include_translations', 'false').lower() == 'true'
    translate_to = request.args.get('to_lang', '')
    
    try:
        # Get article content
        article = wiki_utils.get_article_content(title, lang)
        
        if not article:
            flash('Article not found', 'danger')
            return redirect(url_for('home'))
        
        # Process article content into sections
        sections = wiki_utils.split_content_into_sections(article['content'])
        
        # Get translations if needed
        translations = {}
        if include_translations and translate_to:
            try:
                # Get translation of summary
                translations['summary'] = wiki_utils.translate_text(article['summary'], translate_to, lang)
                
                # Get translations of sections
                translations['sections'] = []
                for section in sections:
                    if section['content']:
                        translated_content = wiki_utils.translate_text(section['content'], translate_to, lang)
                        translations['sections'].append({
                            'title': section['title'],
                            'content': translated_content
                        })
                    else:
                        translations['sections'].append({
                            'title': section['title'],
                            'content': ''
                        })
            except Exception as e:
                logger.error(f"Translation error: {str(e)}")
                # Continue without translations if they fail
                flash(f"Warning: Some translations could not be completed. The document will include partial translations.", 'warning')
        
        # Create Word document
        doc = Document()
        
        # Add title
        doc.add_heading(article['title'], 0)
        
        # Add metadata
        doc.add_paragraph(f"Source: {article['url']}")
        
        if include_translations and translate_to:
            doc.add_paragraph(f"Original Language: {get_language_name(lang)} ({lang})")
            doc.add_paragraph(f"Translated to: {get_language_name(translate_to)} ({translate_to})")
        else:
            doc.add_paragraph(f"Language: {get_language_name(lang)} ({lang})")
        
        doc.add_paragraph(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add summary
        doc.add_heading('Summary', 1)
        doc.add_paragraph(article['summary'])
        
        # Add translated summary if available
        if include_translations and translate_to and 'summary' in translations:
            doc.add_heading(f'Summary (Translated to {get_language_name(translate_to)})', 2)
            doc.add_paragraph(translations['summary'])
        
        # Add content sections
        doc.add_heading('Full Content', 1)
        
        for i, section in enumerate(sections):
            if section['title']:
                level = min(section['level'] + 1, 9)  # Word doc supports headings 1-9
                doc.add_heading(section['title'], level)
            
            # Add original content
            doc.add_paragraph(section['content'])
            
            # Add translation if requested
            if (include_translations and translate_to and 
                'sections' in translations and i < len(translations['sections'])):
                
                # Add a heading for the translation if there's content to translate
                if section['content'].strip():
                    translated_heading = f"Translation ({get_language_name(translate_to)})"
                    doc.add_heading(translated_heading, level + 1 if section['title'] else level)
                    
                    # Add the translated content
                    doc.add_paragraph(translations['sections'][i]['content'])
                    
                    # Add a separator for readability
                    doc.add_paragraph("---")
        
        # Save to memory
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Return file for download
        safe_filename = title.replace('/', '_').replace('\\', '_')
        
        # Create a filename that indicates if translations are included
        if include_translations and translate_to:
            download_name = f"{safe_filename}_{lang}_with_{translate_to}_translations.docx"
        else:
            download_name = f"{safe_filename}_{lang}.docx"
            
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=download_name
        )
    except Exception as e:
        logger.error(f"Export error: {str(e)}")
        flash(f"Error exporting document: {str(e)}", 'danger')
        return redirect(url_for('view_article', title=title, lang=lang))


def keep_alive():
    while True:
        try:
            print("[KeepAlive] Pinging self to stay awake...")
            requests.get("https://wikitruth.onrender.com/")  # ЗАМЕНИ НА СВОЙ URL!
        except Exception as e:
            print("[KeepAlive] Ping failed:", e)
        time.sleep(600)  # каждые 10 минут

# Error handlers
@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

@app.errorhandler(500)
def server_error(e):
    return render_template('500.html'), 500
